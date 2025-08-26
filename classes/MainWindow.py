from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QFormLayout, QLineEdit, QComboBox, 
    QPushButton, QTextEdit, QCheckBox, QFileDialog, QMessageBox, QGroupBox, 
    QHBoxLayout, QLabel, QScrollArea, QDialog, QButtonGroup, QRadioButton
)
from PySide6.QtCore import Qt, QDate
from PySide6.QtPrintSupport import QPrintDialog, QPrinter
from PySide6.QtGui import QTextDocument
from classes.Patient import PatientData, Gender, CYP2C19, ABCB1
from logic.Mod1 import mod1, mod1_text
from logic.Mod2 import mod2
from logic.Mod3 import mod3
from logic.Mod4 import mod4
from logic.Mod5 import mod5
from logic.exel_utils import append_patient_data, DEFAULT_FILENAME, calculate_ckd_epi, calculate_creatinine_clearance
from logic.html_utils import format_html_table, format_html_table_advanced
from logic.word_utils import add_table_with_title
from logic.validation_utils import (
    validate_age, validate_weight, validate_height, validate_creatinine,
    validate_mpv, validate_plcr, validate_spontaneous_aggregation,
    validate_induced_aggregation_1_ADP, validate_induced_aggregation_5_ADP,
    validate_induced_aggregation_15_ARA, validate_platelet_count,
    get_drug_cancellation_recommendation
)
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math

class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Антиагрегантная терапия")
        self.resize(1000, 800)
        
        # Создаем главный layout
        main_layout = QVBoxLayout(self)
        
        # Создаем область прокрутки
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Создаем контейнер для содержимого
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)

        # === ГРУППА 1: ОСНОВНЫЕ ДАННЫЕ ПАЦИЕНТА ===
        basic_group = QGroupBox("Основные данные пациента")
        basic_layout = QFormLayout()
        
        # Поля ввода
        self.date = QLineEdit()
        self.date.setPlaceholderText("Введите дату (дд.мм.гггг)")
        self.date.setText(QDate.currentDate().toString("dd.MM.yyyy"))
        basic_layout.addRow("Дата обследования:", self.date)

        self.name_or_record = QLineEdit()
        self.name_or_record.setPlaceholderText("Введите ФИО или номер истории болезни")
        basic_layout.addRow("ФИО / № истории болезни:", self.name_or_record)

        self.examination_type = QComboBox()
        self.examination_type.addItems(["Стационар", "Амбулаторно"])
        basic_layout.addRow("Обследование:", self.examination_type)
        
        # Поля выбора
        self.gender = QComboBox()
        self.gender.addItem("")  # Для необязательного выбора
        self.gender.addItems([g.value for g in Gender])
        basic_layout.addRow("Пол (выберите):", self.gender)
        
        # Поля ввода
        self.age = QLineEdit()
        self.age.setPlaceholderText("Введите возраст (лет)")
        basic_layout.addRow("Возраст:", self.age)

        self.weight = QLineEdit()
        self.weight.setPlaceholderText("Введите вес (кг)")
        basic_layout.addRow("Вес:", self.weight)

        self.height_field = QLineEdit()
        self.height_field.setPlaceholderText("Введите рост (см)")
        basic_layout.addRow("Рост:", self.height_field)
        
        basic_group.setLayout(basic_layout)
        layout.addWidget(basic_group)

        # === ГРУППА 2: ГЕНОТИПЫ ===
        genotype_group = QGroupBox("Генотипы")
        genotype_layout = QFormLayout()
        
        self.cyp2c19 = QComboBox()
        self.cyp2c19.addItem("")
        self.cyp2c19.addItems([c.value for c in CYP2C19])
        genotype_layout.addRow("Генотип CYP2C19:", self.cyp2c19)

        self.abcb1 = QComboBox()
        self.abcb1.addItem("")
        self.abcb1.addItems(["TT", "TC", "CC"])
        genotype_layout.addRow("Генотип ABCB1:", self.abcb1)
        
        genotype_group.setLayout(genotype_layout)
        layout.addWidget(genotype_group)

        # === ГРУППА 3: БИОХИМИЧЕСКИЕ ПОКАЗАТЕЛИ ===
        bio_group = QGroupBox("Биохимические показатели")
        bio_layout = QFormLayout()
        
        self.creatinine = QLineEdit()
        self.creatinine.setPlaceholderText("Введите креатинин (мкмоль/л)")
        bio_layout.addRow("Креатинин:", self.creatinine)

        self.creatinine_clearance = QLineEdit()
        self.creatinine_clearance.setPlaceholderText("Введите клиренс креатинина (мл/мин)")
        bio_layout.addRow("Клиренс креатинина:", self.creatinine_clearance)
        
        bio_group.setLayout(bio_layout)
        layout.addWidget(bio_group)

        # === ГРУППА 4: ТРОМБОЦИТАРНЫЕ ПОКАЗАТЕЛИ ===
        platelet_group = QGroupBox("Тромбоцитарные показатели")
        platelet_layout = QFormLayout()
        
        # Добавляем поле для количества тромбоцитов
        self.platelet_count = QLineEdit()
        self.platelet_count.setPlaceholderText("Введите количество тромбоцитов (×10⁹/л)")
        platelet_layout.addRow("Количество тромбоцитов, ×10⁹/л:", self.platelet_count)
        
        self.mpv = QLineEdit()
        self.mpv.setPlaceholderText("Введите MPV (фл)")
        platelet_layout.addRow("Величина тромбоцитов MPV:", self.mpv)

        self.plcr = QLineEdit()
        self.plcr.setPlaceholderText("Введите PLCR (%)")
        platelet_layout.addRow("Отн. кол-во больших тромбоцитов PLCR:", self.plcr)
        
        platelet_group.setLayout(platelet_layout)
        layout.addWidget(platelet_group)

        # === ГРУППА 5: АГРЕГАЦИя ТРОМБОЦИТОВ ===
        aggregation_group = QGroupBox("Агрегация тромбоцитов")
        aggregation_layout = QFormLayout()
        
        self.spontaneous_aggregation = QLineEdit()
        self.spontaneous_aggregation.setPlaceholderText("Введите спонтанную агрегацию (усл.ед.)")
        aggregation_layout.addRow("Спонтанная агрегация:", self.spontaneous_aggregation)

        self.induced_aggregation_1_ADP = QLineEdit()
        self.induced_aggregation_1_ADP.setPlaceholderText("Введите % агрегации")
        aggregation_layout.addRow("Индуц. агрегация 1 мкМоль АДФ:", self.induced_aggregation_1_ADP)

        self.induced_aggregation_5_ADP = QLineEdit()
        self.induced_aggregation_5_ADP.setPlaceholderText("Введите % агрегации")
        aggregation_layout.addRow("Индуц. агрегация 5 мкМоль АДФ:", self.induced_aggregation_5_ADP)

        self.induced_aggregation_15_ARA = QLineEdit()
        self.induced_aggregation_15_ARA.setPlaceholderText("Введите % агрегации")
        aggregation_layout.addRow("Индуц. агрегация 15 мкл арахидоновой кислоты:", self.induced_aggregation_15_ARA)
        
        aggregation_group.setLayout(aggregation_layout)
        layout.addWidget(aggregation_group)

        # === ГРУППА 6: ОЦЕНКА РИСКА ЖЕЛУДОЧНО-КИШЕЧНОГО КРОВОТЕЧЕНИЯ ===
        gi_bleeding_group = QGroupBox("Оценка риска желудочно-кишечного кровотечения")
        gi_bleeding_layout = QFormLayout()
        
        # Поля для оценки риска ЖКК
        self.ulcer_history = QComboBox()
        self.ulcer_history.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Язвенная болезнь в анамнезе:", self.ulcer_history)

        self.gi_bleeding_history = QComboBox()
        self.gi_bleeding_history.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Желудочно-кишечное кровотечение в анамнезе:", self.gi_bleeding_history)

        self.nsaid_use = QComboBox()
        self.nsaid_use.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Использование НПВП:", self.nsaid_use)

        self.steroid_use = QComboBox()
        self.steroid_use.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Прием ГКС:", self.steroid_use)

        self.age_65 = QComboBox()
        self.age_65.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Возраст ≥ 65 лет:", self.age_65)

        self.dyspepsia = QComboBox()
        self.dyspepsia.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Диспепсия:", self.dyspepsia)

        self.gerd = QComboBox()
        self.gerd.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Желудочно-пищеводный рефлюкс:", self.gerd)

        self.h_pylori = QComboBox()
        self.h_pylori.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Инфицирование H. pylori:", self.h_pylori)

        self.alcohol_use = QComboBox()
        self.alcohol_use.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Хроническое употребление алкоголя:", self.alcohol_use)
        
        gi_bleeding_group.setLayout(gi_bleeding_layout)
        layout.addWidget(gi_bleeding_group)

        # === ГРУППА 7: ПРЕПАРАТЫ ===
        drugs_group = QGroupBox("Препараты")
        drugs_layout = QVBoxLayout()
        
        drugs_label = QLabel("Выберите принимаемый препарат:")
        drugs_layout.addWidget(drugs_label)
        
        # Создаем группу радиокнопок для выбора только одного препарата
        self.drugs_button_group = QButtonGroup(self)
        
        self.drug_aspirin = QRadioButton("АСК")
        self.drug_clopidogrel = QRadioButton("Клопидогрел")
        self.drug_aspirin_clopidogrel = QRadioButton("АСК+клопидогрел")
        self.drug_aspirin_ticagrelor = QRadioButton("АСК+тикагрелор")
        
        # Добавляем радиокнопки в группу
        self.drugs_button_group.addButton(self.drug_aspirin, 1)
        self.drugs_button_group.addButton(self.drug_clopidogrel, 2)
        self.drugs_button_group.addButton(self.drug_aspirin_clopidogrel, 3)
        self.drugs_button_group.addButton(self.drug_aspirin_ticagrelor, 4)
        
        # Устанавливаем "АСК" по умолчанию
        self.drug_aspirin.setChecked(True)
        
        drugs_layout.addWidget(self.drug_aspirin)
        drugs_layout.addWidget(self.drug_clopidogrel)
        drugs_layout.addWidget(self.drug_aspirin_clopidogrel)
        drugs_layout.addWidget(self.drug_aspirin_ticagrelor)
        
        drugs_group.setLayout(drugs_layout)
        layout.addWidget(drugs_group)

        # === ГРУППА 7: ДЕЙСТВИЯ ===
        actions_group = QGroupBox("⚙️ Действия")
        actions_layout = QVBoxLayout()
        
        self.report_button = QPushButton("📄 Сформировать полный отчет")
        self.report_button.clicked.connect(self.generate_report)
        actions_layout.addWidget(self.report_button)

        # Убрана кнопка сохранения в PDF
        self.save_doc_button = QPushButton("📝 Сохранить отчет в DOC")
        self.save_doc_button.clicked.connect(self.save_report_to_doc)
        actions_layout.addWidget(self.save_doc_button)
        
        actions_group.setLayout(actions_layout)
        layout.addWidget(actions_group)

        # Устанавливаем контейнер в область прокрутки
        scroll_area.setWidget(content_widget)
        
        # Добавляем область прокрутки в главный layout
        main_layout.addWidget(scroll_area)

        self.excel_filename = DEFAULT_FILENAME
        self.patient_data = None
        self.current_report_html = ""
        self.current_report_data = None

        # Применяем стили
        self.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin-top: 1ex;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
            QLineEdit {
                padding: 5px;
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                background-color: #f8f9fa;
            }
            QLineEdit:focus {
                border: 2px solid #3498db;
                background-color: white;
            }
            QComboBox {
                padding: 5px;
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                background-color: #ecf0f1;
            }
            QComboBox:focus {
                border: 2px solid #3498db;
            }
            QPushButton {
                padding: 8px 16px;
                border: 2px solid #3498db;
                border-radius: 5px;
                background-color: #3498db;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
            QRadioButton {
                spacing: 8px;
                font-weight: normal;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
            QScrollArea {
                border: none;
                background-color: transparent;
            }
        """)

        # Подключение валидации к полям
        self.age.textChanged.connect(self.validate_age)
        self.weight.textChanged.connect(self.validate_weight)
        self.height_field.textChanged.connect(self.validate_height)
        self.creatinine.textChanged.connect(self.validate_creatinine)
        self.creatinine_clearance.textChanged.connect(self.validate_creatinine_clearance)
        self.mpv.textChanged.connect(self.validate_mpv)
        self.plcr.textChanged.connect(self.validate_plcr)
        self.spontaneous_aggregation.textChanged.connect(self.validate_spontaneous_aggregation)
        self.induced_aggregation_1_ADP.textChanged.connect(self.validate_induced_aggregation_1_ADP)
        self.induced_aggregation_5_ADP.textChanged.connect(self.validate_induced_aggregation_5_ADP)
        self.induced_aggregation_15_ARA.textChanged.connect(self.validate_induced_aggregation_15_ARA)
        self.platelet_count.textChanged.connect(self.validate_platelet_count)

    # Методы валидации
    def validate_age(self):
        try:
            age = int(self.age.text())
            if age <= 0 or age > 120:
                self.age.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.age.setStyleSheet("")
                return True
        except ValueError:
            self.age.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
            return False

    def validate_weight(self):
        try:
            weight = float(self.weight.text())
            if weight <= 0 or weight > 300:
                self.weight.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.weight.setStyleSheet("")
                return True
        except ValueError:
            self.weight.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
            return False

    def validate_height(self):
        try:
            height = float(self.height_field.text())
            if height <= 0 or height > 250:
                self.height_field.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.height_field.setStyleSheet("")
                return True
        except ValueError:
            if self.height_field.text():
                self.height_field.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.height_field.setStyleSheet("")
                return True

    def validate_creatinine(self):
        try:
            creatinine = float(self.creatinine.text())
            if creatinine <= 0 or creatinine > 1000:
                self.creatinine.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.creatinine.setStyleSheet("")
                return True
        except ValueError:
            if self.creatinine.text():
                self.creatinine.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.creatinine.setStyleSheet("")
                return True

    def validate_creatinine_clearance(self):
        try:
            clearance = float(self.creatinine_clearance.text())
            if clearance <= 0 or clearance > 200:
                self.creatinine_clearance.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.creatinine_clearance.setStyleSheet("")
                return True
        except ValueError:
            if self.creatinine_clearance.text():
                self.creatinine_clearance.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.creatinine_clearance.setStyleSheet("")
                return True

    def validate_mpv(self):
        try:
            mpv = float(self.mpv.text())
            if mpv <= 0 or mpv > 20:
                self.mpv.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.mpv.setStyleSheet("")
                return True
        except ValueError:
            if self.mpv.text():
                self.mpv.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.mpv.setStyleSheet("")
                return True

    def validate_plcr(self):
        try:
            plcr = float(self.plcr.text())
            if plcr < 0 or plcr > 100:
                self.plcr.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.plcr.setStyleSheet("")
                return True
        except ValueError:
            if self.plcr.text():
                self.plcr.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.plcr.setStyleSheet("")
                return True

    def validate_spontaneous_aggregation(self):
        try:
            agg = float(self.spontaneous_aggregation.text())
            if agg < 0 or agg > 100:
                self.spontaneous_aggregation.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.spontaneous_aggregation.setStyleSheet("")
                return True
        except ValueError:
            if self.spontaneous_aggregation.text():
                self.spontaneous_aggregation.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.spontaneous_aggregation.setStyleSheet("")
                return True

    def validate_induced_aggregation_1_ADP(self):
        try:
            agg = float(self.induced_aggregation_1_ADP.text())
            if agg < 0 or agg > 100:
                self.induced_aggregation_1_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_1_ADP.setStyleSheet("")
                return True
        except ValueError:
            if self.induced_aggregation_1_ADP.text():
                self.induced_aggregation_1_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_1_ADP.setStyleSheet("")
                return True

    def validate_induced_aggregation_5_ADP(self):
        try:
            agg = float(self.induced_aggregation_5_ADP.text())
            if agg < 0 or agg > 100:
                self.induced_aggregation_5_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_5_ADP.setStyleSheet("")
                return True
        except ValueError:
            if self.induced_aggregation_5_ADP.text():
                self.induced_aggregation_5_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_5_ADP.setStyleSheet("")
                return True

    def validate_induced_aggregation_15_ARA(self):
        try:
            agg = float(self.induced_aggregation_15_ARA.text())
            if agg < 0 or agg > 100:
                self.induced_aggregation_15_ARA.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_15_ARA.setStyleSheet("")
                return True
        except ValueError:
            if self.induced_aggregation_15_ARA.text():
                self.induced_aggregation_15_ARA.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_15_ARA.setStyleSheet("")
                return True

    def validate_platelet_count(self):
        try:
            platelets = float(self.platelet_count.text())
            if platelets <= 0 or platelets > 1000:
                self.platelet_count.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.platelet_count.setStyleSheet("")
                return True
        except ValueError:
            if self.platelet_count.text():
                self.platelet_count.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.platelet_count.setStyleSheet("")
                return True

    def validate_all_fields(self):
        validations = [
            self.validate_age(),
            self.validate_weight(),
            self.validate_height(),
            self.validate_creatinine(),
            self.validate_creatinine_clearance(),
            self.validate_platelet_count(),
            self.validate_mpv(),
            self.validate_plcr(),
            self.validate_spontaneous_aggregation(),
            self.validate_induced_aggregation_1_ADP(),
            self.validate_induced_aggregation_5_ADP(),
            self.validate_induced_aggregation_15_ARA()
        ]

        if not all(validations):
            QMessageBox.warning(self, "Ошибка валидации",
                              "Пожалуйста, исправьте ошибки в полях (выделены красным)")
            return False
        return True

    def get_selected_drug(self):
        if self.drug_aspirin.isChecked():
            return "АСК"
        elif self.drug_clopidogrel.isChecked():
            return "клопидогрел"
        elif self.drug_aspirin_clopidogrel.isChecked():
            return "АСК+клопидогрел"
        elif self.drug_aspirin_ticagrelor.isChecked():
            return "АСК+тикагрелор"
        else:
            return ""

    def calculate_gi_bleeding_score(self):
        """Рассчитывает сумму баллов для оценки риска ЖКК"""
        score = 0
        
        # Преобразуем "да"/"нет" в 1/0
        if self.ulcer_history.currentText() == "да":
            score += 1
        if self.gi_bleeding_history.currentText() == "да":
            score += 1
        if self.nsaid_use.currentText() == "да":
            score += 1
        if self.steroid_use.currentText() == "да":
            score += 1
        if self.age_65.currentText() == "да":
            score += 1
        if self.dyspepsia.currentText() == "да":
            score += 1
        if self.gerd.currentText() == "да":
            score += 1
        if self.h_pylori.currentText() == "да":
            score += 1
        if self.alcohol_use.currentText() == "да":
            score += 1
            
        return score

    def format_html_table(self, headers, rows):
        html = '<table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 12px;">'
        
        # Заголовки
        html += '<tr style="background-color: #f2f2f2; font-weight: bold;">'
        for header in headers:
            html += f'<th style="border: 1px solid #000; padding: 8px; text-align: center;">{header}</th>'
        html += '</tr>'
        
        # Данные
        for row in rows:
            html += '<tr>'
            for cell in row:
                html += f'<td style="border: 1px solid #000; padding: 8px; text-align: center;">{cell}</td>'
            html += '</tr>'
        
        html += '</table>'
        return html

    def save_report_to_doc(self):
        if not hasattr(self, 'current_report_data') or not self.current_report_data:
            QMessageBox.warning(self, "Предупреждение", "Сначала сформируйте отчет")
            return
            
        filename, _ = QFileDialog.getSaveFileName(
            self, 
            "Сохранить отчет в DOC", 
            f"медицинский_отчет_{QDate.currentDate().toString('yyyy-MM-dd')}.docx",
            "Word Documents (*.docx);;Все файлы (*)"
        )
        if filename:
            try:
                # Создаем новый документ
                doc = docx.Document()
                
                # Добавляем заголовок
                title = doc.add_heading('РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем основную информацию
                doc.add_paragraph(f"Дата обследования: {self.current_report_data['date']}")
                doc.add_paragraph(f"ФИО: {self.current_report_data['name']}")
                doc.add_paragraph(f"Возраст: {self.current_report_data['age']}")
                doc.add_paragraph()
                
                # Добавляем информацию о препаратах
                doc.add_paragraph().add_run("Прием антиагрегантов:").bold = True
                doc.add_paragraph(f"Антиагреганты, которые пациент принимает: {self.current_report_data['drugs']}")
                doc.add_paragraph()
                
                # Добавляем таблицы с данными
                add_table_with_title(doc, 
                    ["Параметр", "Результат пациента", "Критерий", "Оценка", "Прогноз"],
                    self.current_report_data['main_table_rows'],
                    "Прием антиагрегатов:"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента", "Оценка метаболизма", "Рекомендации"],
                    self.current_report_data['cyp_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ ГЕНОТИПА CYP 2C19"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента ABCB1", "Оценка транспорта", "Рекомендации"],
                    self.current_report_data['abcb1_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ АКТИВНОСТИ ТРАНСПОРТНОЙ СИСТЕМЫ P-ГЛИКОПРОТЕИНА"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['ticagrelor_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ ТИКАГРЕЛОРОМ"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 15 мкл арахидоновой кислоты, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['aspirin_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ АЦЕТИЛСАЛИЦИЛОВОЙ КИСЛОТОЙ"
                )
                
                # Сохраняем документ
                doc.save(filename)
                QMessageBox.information(self, "Сохранение", f"Отчет сохранен в DOC файл:\n{filename}")
                
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить DOC файл:\n{str(e)}")
                print(f"Ошибка при сохранении DOC: {e}")

    def generate_report(self):
        try:
            if not self.validate_all_fields():
                return
            
            # Сбор данных
            date = self.date.text() if self.date.text() else QDate.currentDate().toString("dd.MM.yyyy")
            name_or_record = self.name_or_record.text() if self.name_or_record.text() else "____________________________________"
            age = int(self.age.text()) if self.age.text() else 0
            examination_type = self.examination_type.currentText()
            gender = self.gender.currentText()
            weight = float(self.weight.text()) if self.weight.text() else 0
            creatinine = float(self.creatinine.text()) if self.creatinine.text() else 0
            
            # Расчет КК и СКФ
            ccr = calculate_creatinine_clearance(age, weight, gender, creatinine)
            gfr = calculate_ckd_epi(age, gender, creatinine)
            
            # Получаем количество тромбоцитов
            platelet_count = self.platelet_count.text() if self.platelet_count.text() else "______"
            
            # Получаем рекомендации по отмене препаратов
            selected_drug = self.get_selected_drug()
            drug_cancellation = get_drug_cancellation_recommendation(platelet_count, selected_drug)
            
            # Расчет оценки риска ЖКК
            gi_bleeding_score = self.calculate_gi_bleeding_score()
            
            # Получаем данные агрегации
            T_adp = float(self.induced_aggregation_5_ADP.text()) if self.induced_aggregation_5_ADP.text() else None
            T_ara = float(self.induced_aggregation_15_ARA.text()) if self.induced_aggregation_15_ARA.text() else None
            
            # Генетические данные
            cyp_genotype = self.cyp2c19.currentText() if self.cyp2c19.currentText() else "______"
            abcb1_genotype = self.abcb1.currentText() if self.abcb1.currentText() else "______"
            
            # Данные о терапии
            selected_drug = self.get_selected_drug()
            drugs_str = selected_drug if selected_drug else "___________"

            # Расчет коэффициента прогноза
            try:
                gender_val = 1 if self.gender.currentText() == "Муж" else 2 if self.gender.currentText() == "Жен" else 0
                prognosis_value = mod1(
                    gender_val,
                    float(self.age.text()) if self.age.text() else 0,
                    float(self.weight.text()) if self.weight.text() else 0,
                    float(self.height_field.text()) if self.height_field.text() else 0,
                    float(self.creatinine.text()) if self.creatinine.text() else 0,
                    float(self.creatinine_clearance.text()) if self.creatinine_clearance.text() else 0,
                    float(self.mpv.text()) if self.mpv.text() else 0,
                    float(self.plcr.text()) if self.plcr.text() else 0,
                    float(self.spontaneous_aggregation.text()) if self.spontaneous_aggregation.text() else 0,
                    float(self.induced_aggregation_1_ADP.text()) if self.induced_aggregation_1_ADP.text() else 0,
                    float(self.induced_aggregation_5_ADP.text()) if self.induced_aggregation_5_ADP.text() else 0,
                    float(self.induced_aggregation_15_ARA.text()) if self.induced_aggregation_15_ARA.text() else 0
                )
                prognosis_evaluation = mod1_text(prognosis_value)
            except Exception as e:
                prognosis_value = "Ошибка расчета"
                prognosis_evaluation = ("Ошибка", ["Ошибка расчета коэффициента прогноза"])

            # Сохраняем данные для DOC экспорта
            self.current_report_data = {
                'date': date,
                'name_or_record': name_or_record,
                'examination_type': examination_type,
                'age': age,
                'drugs': drugs_str,
                'main_table_rows': [],
                'cyp_table_rows': [],
                'abcb1_table_rows': [],
                'ticagrelor_table_rows': [],
                'aspirin_table_rows': [],
                'gi_bleeding_table_rows': []
            }

            # Формируем HTML отчет и данные для таблиц
            html_report = f"""
            <html>
            <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ text-align: center; font-size: 16px; font-weight: bold; margin-bottom: 20px; }}
                .section {{ margin: 20px 0; }}
                .section-title {{ font-size: 14px; font-weight: bold; margin-bottom: 10px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 8px; text-align: center; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
            </style>
            </head>
            <body>
                <div class="header">РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ</div>
                
                <p><strong>Дата обследования:</strong> {date}</p>
                <p><strong>ФИО / № истории болезни:</strong> {name_or_record}</p>
                <p><strong>Обследование:</strong> {examination_type}</p>
                <p><strong>Возраст:</strong> {age}</p>
                
                <div class="section">
                    <div class="section-title">Прием антиагрегантов:</div>
                    <p><strong>Антиагреганты, которые пациент принимает:</strong> {drugs_str}</p>
                </div>
            """

            # Таблица 1: Основные результаты
            main_table_headers = ["Параметр", "Результат пациента", "Критерий", "Оценка", "Прогноз"]
            main_table_rows = []

            # Строка 1: Коэффициент прогноза
            if isinstance(prognosis_value, (int, float)):
                if prognosis_value <= 1.56:
                    criterion = "≤ 1.56"
                    evaluation = "Благоприятная"
                    prognosis_text = "Неблагоприятных событий в течение года не ожидается"
                elif 1.561 <= prognosis_value <= 2.087:
                    criterion = "1.561-2.087"
                    evaluation = "Неблагоприятная"
                    prognosis_text = "Возможны обращения за медицинской помощью в течение ближайшего года"
                else:
                    criterion = "> 2.08"
                    evaluation = "Риск повторных сосудистых событий"
                    prognosis_text = "Высокий риск повторного инфаркта и летальный исход"
                main_table_rows.append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    f"{prognosis_value:.3f}",
                    criterion,
                    evaluation,
                    prognosis_text
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    f"{prognosis_value:.3f}",
                    criterion,
                    evaluation,
                    prognosis_text
                ])
            else:
                main_table_rows.append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    prognosis_value,
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    prognosis_value,
                    "-",
                    "-",
                    "-"
                ])

            # Строка 2: Индуцированная агрегация 5 мкМоль АДФ
            if T_adp is not None:
                if T_adp <= 10:
                    criterion_adp = "T ≤ 10 %"
                    evaluation_adp = "Агрегация тромбоцитов значительно подавлена"
                    prognosis_adp = "Риск геморрагических осложнений"
                elif 10 < T_adp < 25:
                    criterion_adp = "10 < T < 25 %"
                    evaluation_adp = "Агрегация тромбоцитов умеренно подавлена"
                    prognosis_adp = "Терапия эффективна"
                else:
                    criterion_adp = "T ≥ 25 %"
                    evaluation_adp = "Агрегация тромбоцитов сохранена"
                    prognosis_adp = "Терапия неэффективна"

                main_table_rows.append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    f"{T_adp}%",
                    criterion_adp,
                    evaluation_adp,
                    prognosis_adp
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    f"{T_adp}%",
                    criterion_adp,
                    evaluation_adp,
                    prognosis_adp
                ])
            else:
                main_table_rows.append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    "______",
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    "______",
                    "-",
                    "-",
                    "-"
                ])

            # Строка 3: Генотип CYP 2C19
            if cyp_genotype != "______":
                if cyp_genotype == "CYP 2c19*1":
                    evaluation_cyp = "Нормальный метаболизм клопидогрела"
                    prognosis_cyp = "Эффективность терапии клопидогрелом"
                elif cyp_genotype in ["CYP 2c19*2", "CYP 2c19*3"]:
                    evaluation_cyp = "Замедленный метаболизм клопидогрела"
                    prognosis_cyp = "Возможна резистентность к клопидогрелу"
                elif cyp_genotype == "CYP 2c19*17":
                    evaluation_cyp = "Ускоренный метаболизм клопидогрела"
                    prognosis_cyp = "Возможно угнетение агрегации, риск геморрагических осложнений"
                else:
                    evaluation_cyp = "Неизвестный генотип"
                    prognosis_cyp = "Требуется дополнительное исследование"

                main_table_rows.append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    cyp_genotype,
                    cyp_genotype,
                    evaluation_cyp,
                    prognosis_cyp
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    cyp_genotype,
                    cyp_genotype,
                    evaluation_cyp,
                    prognosis_cyp
                ])
            else:
                main_table_rows.append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    "______",
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    "______",
                    "-",
                    "-",
                    "-"
                ])

            # Строка 4: Генотип ABCB1
            if abcb1_genotype != "______":
                if abcb1_genotype == "TT":
                    evaluation_abcb1 = "Выведение клопидогрела ускорено"
                    prognosis_abcb1 = "Вероятна резистентность к клопидогрелу"
                elif abcb1_genotype == "TC":
                    evaluation_abcb1 = "Незначительное ускорение выведения клопидогрела"
                    prognosis_abcb1 = "Клинически незначимое влияние эффективность фармакотерапии"
                elif abcb1_genotype == "CC":
                    evaluation_abcb1 = "Нормальная скорость выведения клопидогрела"
                    prognosis_abcb1 = "Влияния на эффективность терапии клопидогрелом нет"
                else:
                    evaluation_abcb1 = "Неизвестный генотип"
                    prognosis_abcb1 = "Требуется дополнительное исследование"

                main_table_rows.append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    abcb1_genotype,
                    abcb1_genotype,
                    evaluation_abcb1,
                    prognosis_abcb1
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    abcb1_genotype,
                    abcb1_genotype,
                    evaluation_abcb1,
                    prognosis_abcb1
                ])
            else:
                main_table_rows.append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    "______",
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    "______",
                    "-",
                    "-",
                    "-"
                ])

            # Добавляем основную таблицу в отчет
            html_report += self.format_html_table(main_table_headers, main_table_rows)

            # Таблица 2: Коррекция терапии клопидогрелом (CYP2C19)
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ ГЕНОТИПА CYP 2C19</div>
            """

            cyp_table_headers = [
                "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Генотип пациента",
                "Оценка метаболизма",
                "Рекомендации"
            ]
            
            cyp_table_rows = []
            if T_adp is not None and cyp_genotype != "______":
                if T_adp <= 10:
                    criterion = "T ≤ 10 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                elif 10 < T_adp < 25:
                    criterion = "10 < T < 25 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                else:
                    criterion = "T ≥ 25 %"
                    state = "Агрегация тромбоцитов сохранена"
                
                if cyp_genotype == "CYP 2c19*1":
                    metabolism = "Нормальный метаболизм"
                    if T_adp <= 10:
                        recommendation = "Продолжить прием клопидогрела. Риск геморрагических осложнений"
                    elif 10 < T_adp < 25:
                        recommendation = "Продолжить прием клопидогрела. Терапия эффективна"
                    else:
                        recommendation = "Определить комплаентность пациента. Контроль агрегации через 5 дней"
                elif cyp_genotype in ["CYP 2c19*2", "CYP 2c19*3"]:
                    metabolism = "Замедленный метаболизм"
                    if T_adp <= 10:
                        recommendation = "Продолжить прием клопидогрела. Риск геморрагических осложнений"
                    elif 10 < T_adp < 25:
                        recommendation = "Продолжить прием клопидогрела. Терапия эффективна"
                    else:
                        recommendation = "Замена на прасугрел или тикагрелор. Контроль агрегации через 5 дней"
                elif cyp_genotype == "CYP 2c19*17":
                    metabolism = "Ускоренный метаболизм"
                    if T_adp <= 10:
                        recommendation = "Снизить дозу клопидогрела. Высокий риск геморрагических осложнений"
                    elif 10 < T_adp < 25:
                        recommendation = "Продолжить прием клопидогрела. Терапия эффективна"
                    else:
                        recommendation = "Определить комплаентность пациента. Контроль агрегации через 5 дней"
                else:
                    metabolism = "Неизвестный метаболизм"
                    recommendation = "Требуется дополнительное исследование"
                
                cyp_table_rows.append([f"{T_adp}%", criterion, state, cyp_genotype, metabolism, recommendation])
                self.current_report_data['cyp_table_rows'].append([f"{T_adp}%", criterion, state, cyp_genotype, metabolism, recommendation])
            else:
                cyp_table_rows.append(["______", "-", "-", "-", "-", "-"])
                self.current_report_data['cyp_table_rows'].append(["______", "-", "-", "-", "-", "-"])

            html_report += self.format_html_table(cyp_table_headers, cyp_table_rows)
            html_report += "</div>"

            # Таблица 3: Коррекция терапии клопидогрелом (ABCB1)
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ АКТИВНОСТИ ТРАНСПОРТНОЙ СИСТЕМЫ Р-ГЛИКОПРОТЕИНА</div>
            """

            abcb1_table_headers = [
                "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Генотип пациента ABCB1",
                "Оценка транспорта",
                "Рекомендации"
            ]
            
            abcb1_table_rows = []
            if T_adp is not None and abcb1_genotype != "______":
                if T_adp <= 10:
                    criterion = "T ≤ 10 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                elif 10 < T_adp < 25:
                    criterion = "10 < T < 25 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                else:
                    criterion = "T ≥ 25 %"
                    state = "Агрегация тромбоцитов сохранена"
                
                if abcb1_genotype == "TT":
                    transport = "Ускоренное выведение"
                    if T_adp <= 10:
                        recommendation = "Продолжить прием клопидогрела. Риск геморрагических осложнений"
                    elif 10 < T_adp < 25:
                        recommendation = "Продолжить прием клопидогрела. Терапия эффективна"
                    else:
                        recommendation = "Увеличить дозу клопидогрела или замена на другой антиагрегант"
                elif abcb1_genotype == "TC":
                    transport = "Незначительно ускоренное выведение"
                    if T_adp <= 10:
                        recommendation = "Продолжить прием клопидогрела. Риск геморрагических осложнений"
                    elif 10 < T_adp < 25:
                        recommendation = "Продолжить прием клопидогreла. Терапия эффективна"
                    else:
                        recommendation = "Контроль агрегации через 5 дней. Рассмотреть увеличение дозы"
                elif abcb1_genotype == "CC":
                    transport = "Нормальное выведение"
                    if T_adp <= 10:
                        recommendation = "Продолжить прием клопидогрела. Риск геморрагических осложнений"
                    elif 10 < T_adp < 25:
                        recommendation = "Продолжить прием клопидогрела. Терапия эффективна"
                    else:
                        recommendation = "Определить комплаентность пациента. Контроль агрегации через 5 дней"
                else:
                    transport = "Неизвестный транспорт"
                    recommendation = "Требуется дополнительное исследование"
                
                abcb1_table_rows.append([f"{T_adp}%", criterion, state, abcb1_genotype, transport, recommendation])
                self.current_report_data['abcb1_table_rows'].append([f"{T_adp}%", criterion, state, abcb1_genotype, transport, recommendation])
            else:
                abcb1_table_rows.append(["______", "-", "-", "-", "-", "-"])
                self.current_report_data['abcb1_table_rows'].append(["______", "-", "-", "-", "-", "-"])

            html_report += self.format_html_table(abcb1_table_headers, abcb1_table_rows)
            html_report += "</div>"

            # Таблица 4: Коррекция терапии тикагрелором
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ ТИКАГРЕЛОРОМ</div>
            """

            ticagrelor_table_headers = [
                "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Рекомендации"
            ]
            
            ticagrelor_table_rows = []
            if T_adp is not None:
                if T_adp <= 10:
                    criterion = "T ≤ 10 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                    recommendation = "Высокий риск геморрагических осложнений. Рассмотреть снижение дозы"
                elif 10 < T_adp < 25:
                    criterion = "10 < T < 25 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                    recommendation = "Продолжить прием тикагрелора. Терапия эффективна"
                else:
                    criterion = "T ≥ 25 %"
                    state = "Агрегация тромбоцитов сохранена"
                    recommendation = "Терапия неэффективна. Замена на другой антиагрегант"
                
                ticagrelor_table_rows.append([f"{T_adp}%", criterion, state, recommendation])
                self.current_report_data['ticagrelor_table_rows'].append([f"{T_adp}%", criterion, state, recommendation])
            else:
                ticagrelor_table_rows.append(["______", "-", "-", "-"])
                self.current_report_data['ticagrelor_table_rows'].append(["______", "-", "-", "-"])

            html_report += self.format_html_table(ticagrelor_table_headers, ticagrelor_table_rows)
            html_report += "</div>"

            # Таблица 5: Коррекция терапии ацетилсалициловой кислотой
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ АЦЕТИЛСАЛИЦИЛОВОЙ КИСЛОТОЙ</div>
            """

            aspirin_table_headers = [
                "Индуцированная агрегация 15 мкл арахидоновой кислоты, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Рекомендации"
            ]
            
            aspirin_table_rows = []
            if T_ara is not None:
                if T_ara <= 2:
                    criterion = "Т ≤ 2 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                    recommendation = "Высокий риск геморрагических осложнений. Продолжить прием ацетилсалициловой кислоты"
                elif 2 < T_ara < 8:
                    criterion = "2 < Т < 8 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                    recommendation = "Продолжить прием ацетилсалициловой кислоты. Риск геморрагических осложнений"
                else:
                    criterion = "Т ≥ 8 %"
                    state = "Агрегация тромбоцитов сохранена"
                    recommendation = "Определить комплаентность пациента. Замена на препарат ацетилсалициловой кислоты другого производителя. Контроль агрегации через 5 дней"
                
                aspirin_table_rows.append([f"{T_ara}%", criterion, state, recommendation])
                self.current_report_data['aspirin_table_rows'].append([f"{T_ara}%", criterion, state, recommendation])
            else:
                aspirin_table_rows.append(["______", "-", "-", "-"])
                self.current_report_data['aspirin_table_rows'].append(["______", "-", "-", "-"])

            html_report += self.format_html_table(aspirin_table_headers, aspirin_table_rows)
            html_report += "</div>"

            # Закрываем HTML
            html_report += """
            </body>
            </html>
            """

            # Сохраняем текущий отчет
            self.current_report_html = html_report

            # Показываем отчет в диалоговом окне
            report_dialog = QDialog(self)
            report_dialog.setWindowTitle("Медицинский отчет")
            report_dialog.resize(900, 700)

            layout = QVBoxLayout(report_dialog)
            text_edit = QTextEdit()
            text_edit.setHtml(html_report)
            text_edit.setReadOnly(True)
            layout.addWidget(text_edit)

            close_button = QPushButton("Закрыть")
            close_button.clicked.connect(report_dialog.accept)
            layout.addWidget(close_button)

            report_dialog.exec()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка генерации отчета", 
                            f"Произошла ошибка при формировании отчета:\n{str(e)}")
            print(f"Ошибка в generate_report: {e}")
            import traceback
            traceback.print_exc()
