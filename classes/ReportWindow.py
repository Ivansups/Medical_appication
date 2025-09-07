from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QTextEdit, QPushButton, 
    QScrollArea, QFileDialog, QMessageBox, QApplication
)
from PySide6.QtCore import Qt, QDate
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from logic.word_utils import add_table_with_title

class ReportWindow(QWidget):
    def __init__(self, report_text, patient_data=None, excel_filename="patients.xlsx"):
        super().__init__()
        self.setWindowTitle("📋 Полный медицинский отчет по пациенту")
        self.resize(900, 700)
        self.patient_data = patient_data
        self.excel_filename = excel_filename
        self.current_report_data = None
        
        # Главный layout
        main_layout = QVBoxLayout(self)
        
        # Заголовок
        header_label = QLabel("МЕДИЦИНСКИЙ ОТЧЕТ ПО ПАЦИЕНТУ")
        header_label.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #2c3e50;
                background-color: #ecf0f1;
                padding: 15px;
                border: 2px solid #3498db;
                border-radius: 8px;
                margin: 5px;
            }
        """)
        header_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(header_label)
        
        # Область прокрутки для текста
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Текстовое поле
        self.text = QTextEdit()
        self.text.setReadOnly(True)
        self.text.setStyleSheet("""
            QTextEdit {
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                font-size: 12px;
                line-height: 1.4;
                background-color: #f8f9fa;
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                padding: 15px;
                color: #2c3e50;
            }
            QTextEdit:focus {
                border: 2px solid #3498db;
            }
        """)
        
        # Форматируем текст отчета
        formatted_text = self.format_report_text(report_text)
        self.text.setText(formatted_text)
        
        scroll_area.setWidget(self.text)
        main_layout.addWidget(scroll_area)
        
        # Кнопки действий
        button_layout = QHBoxLayout()
        
        # Кнопка копирования
        copy_button = QPushButton("📋 Копировать отчет")
        copy_button.clicked.connect(self.copy_to_clipboard)
        copy_button.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                border: 2px solid #27ae60;
                border-radius: 5px;
                background-color: #27ae60;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:pressed {
                background-color: #1e8449;
            }
        """)
        
        # Кнопка сохранения в DOC
        doc_button = QPushButton("📝 Сохранить в DOC")
        doc_button.clicked.connect(self.save_to_doc)
        doc_button.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                border: 2px solid #3498db;
                border-radius: 5px;
                background-color: #3498db;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
        """)
        
        # Кнопка закрытия
        close_button = QPushButton("❌ Закрыть")
        close_button.clicked.connect(self.close)
        close_button.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                border: 2px solid #7f8c8d;
                border-radius: 5px;
                background-color: #7f8c8d;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #636e72;
            }
            QPushButton:pressed {
                background-color: #2d3436;
            }
        """)
        
        button_layout.addWidget(copy_button)
        button_layout.addWidget(doc_button)
        button_layout.addStretch()
        button_layout.addWidget(close_button)
        
        main_layout.addLayout(button_layout)
        
        # Стили для окна
        self.setStyleSheet("""
            QWidget {
                background-color: #ecf0f1;
            }
            QScrollArea {
                border: none;
                background-color: transparent;
            }
            QScrollBar:vertical {
                background-color: #f0f0f0;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #c0c0c0;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #a0a0a0;
            }
        """)
    
    def format_report_text(self, text):
        """Для HTML контента просто возвращаем как есть"""
        return text
    
    def copy_to_clipboard(self):
        """Копирует отчет в буфер обмена"""
        clipboard = QApplication.clipboard()
        clipboard.setText(self.text.toPlainText())
        QMessageBox.information(self, "Копирование", "Отчет скопирован в буфер обмена!")
    
    def save_to_doc(self):
        """Сохраняет отчет в DOC файл с табличным форматом"""
        if not self.current_report_data:
            QMessageBox.warning(self, "Предупреждение", "Нет данных для сохранения в DOC")
            return
            
        filename, _ = QFileDialog.getSaveFileName(
            self, 
            "Сохранить отчет в DOC", 
            f"медицинский_отчет_{QDate.currentDate().toString('yyyy-MM-dd')}.docx",
            "Word Documents (*.docx);;Все файлы (*)"
        )
        if filename:
            try:
                # Создаем новый документ
                doc = docx.Document()
                
                # Добавляем заголовок
                title = doc.add_heading('РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем основную информацию
                doc.add_paragraph(f"Дата обследования: {self.current_report_data['date']}")
                doc.add_paragraph(f"ФИО: {self.current_report_data['name_or_record']}")
                doc.add_paragraph(f"Возраст: {self.current_report_data['age']}")
                doc.add_paragraph()
                
                # Добавляем информацию о препаратах
                doc.add_paragraph().add_run("Прием антиагрегантов:").bold = True
                doc.add_paragraph(f"Антиагреганты, которые пациент принимает: {self.current_report_data['drugs']}")
                doc.add_paragraph()
                
                # Добавляем таблицы с данными
                add_table_with_title(doc, 
                    ["Параметр", "Результат пациента", "Критерий", "Оценка", "Прогноз"],
                    self.current_report_data['main_table_rows'],
                    "Прием антиагрегатов:"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента", "Оценка метаболизма", "Рекомендации"],
                    self.current_report_data['cyp_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ ГЕНОТИПА CYP 2C19"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента ABCB1", "Оценка транспорта", "Рекомендации"],
                    self.current_report_data['abcb1_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ АКТИВНОСТИ ТРАНСПОРТНОЙ СИСТЕМЫ P-ГЛИКОПРОТЕИНА"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['ticagrelor_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ ТИКАГРЕЛОРОМ"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 15 мкл арахидоновой кислоты, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['aspirin_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ АЦЕТИЛСАЛИЦИЛОВОЙ КИСЛОТОЙ"
                )
                
                # Сохраняем документ
                doc.save(filename)
                QMessageBox.information(self, "Сохранение", f"Отчет сохранен в DOC файл:\n{filename}")
                
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить DOC файл:\n{str(e)}")
                print(f"Ошибка при сохранении DOC: {e}")

