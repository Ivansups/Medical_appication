from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_cell_border(cell, border_style="single", border_size=4, border_color="000000"):
    """Устанавливает границы для ячейки таблицы в docx"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Создаем элемент границы
    tcBorders = OxmlElement("w:tcBorders")

    # Устанавливаем границы для всех сторон
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), border_style)
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:color"), border_color)
        tcBorders.append(border)

    tcPr.append(tcBorders)


def add_table_with_title(doc, headers, rows, title, column_widths=None):
    """Добавляет таблицу с заголовком, обеспечивая перенос на новую страницу при необходимости"""
    # Добавляем заголовок таблицы
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].bold = True
    title_para.runs[0].font.size = Pt(12)

    # Добавляем разрыв страницы перед таблицей, если она не помещается
    # Word автоматически перенесет таблицу на следующую страницу при необходимости
    doc.add_paragraph()  # Пустая строка

    # Создаем таблицу
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Устанавливаем ширины колонок если указаны
    if column_widths:
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    # Заполняем заголовки
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)

    # Заполняем данные
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_data)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell)

    # Добавляем пустую строку после таблицы
    doc.add_paragraph()


def add_simple_table(doc, headers, rows, column_widths=None):
    """Добавляет таблицу без заголовка"""
    # Создаем таблицу
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Устанавливаем ширины колонок если указаны
    if column_widths:
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    # Заполняем заголовки
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)

    # Заполняем данные
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_data)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell)

    # Добавляем пустую строку после таблицы
    doc.add_paragraph()
