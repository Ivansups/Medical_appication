from PySide6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QFormLayout, QLineEdit, QComboBox, 
    QPushButton, QTextEdit, QCheckBox, QFileDialog, QMessageBox, QGroupBox, 
    QHBoxLayout, QLabel, QScrollArea, QDialog, QButtonGroup, QRadioButton
)
from PySide6.QtCore import Qt, QDate
from PySide6.QtPrintSupport import QPrintDialog, QPrinter
from PySide6.QtGui import QTextDocument
import sys
from logic.Class import PatientData, Gender, CYP2C19
from logic.Mod1 import mod1, mod1_text
from logic.Mod2 import mod2
from logic.Mod3 import mod3
from logic.Mod4 import mod4
from logic.Mod5 import mod5
import os
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
import math

DEFAULT_FILENAME = "patients.xlsx"

def format_html_table(headers, rows):
    """Форматирует данные в HTML таблицу"""
    html = '<table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse; width: 100%; margin: 10px 0;">'
    
    # Заголовки
    html += '<tr style="background-color: #f2f2f2; font-weight: bold;">'
    for header in headers:
        html += f'<th style="border: 1px solid #ddd; padding: 8px; text-align: center;">{header}</th>'
    html += '</tr>'
    
    # Данные
    for row in rows:
        html += '<tr>'
        for cell in row:
            html += f'<td style="border: 1px solid #ddd; padding: 8px;">{cell}</td>'
        html += '</tr>'
    
    html += '</table>'
    return html

def calculate_ckd_epi(age, gender, creatinine):
    """Расчет СКФ по формуле CKD-EPI"""
    try:
        # Конвертируем креатинин из мкмоль/л в мг/дл (1 мг/дл = 88.4 мкмоль/л)
        scr_mg_dl = float(creatinine) / 88.4
        
        if gender == "Муж":
            k = 0.9
            alpha = -0.302
            gender_factor = 1.0
        else:
            k = 0.7
            alpha = -0.241
            gender_factor = 1.012
        
        scr_k = scr_mg_dl / k
        
        min_val = min(scr_k, 1)
        max_val = max(scr_k, 1)
        
        gfr = 142 * math.pow(min_val, alpha) * math.pow(max_val, -1.2) * math.pow(0.9938, age) * gender_factor
        return round(gfr)
    except:
        return None

def calculate_creatinine_clearance(age, weight, gender, creatinine):
    """Расчет клиренса креатинина по Кокрофту-Голту"""
    try:
        age = float(age)
        weight = float(weight)
        creatinine = float(creatinine)
        
        # Формула Кокрофта-Голта
        if gender == "Муж":
            ccr = ((140 - age) * weight) / (72 * creatinine / 88.4)
        else:
            ccr = ((140 - age) * weight) / (72 * creatinine / 88.4) * 0.85
        
        return round(ccr)
    except:
        return None

def create_or_load_workbook(filename=DEFAULT_FILENAME):
    if os.path.exists(filename):
        wb = load_workbook(filename)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        if ws is None:
            ws = wb.create_sheet("Sheet1")
        headers = [
            "Дата обследования", "ФИО / № истории болезни", "Обследование",
            "Пол", "Возраст", "Вес", "Рост", "Креатинин", "Клиренс креатинина", "СКФ",
            "Количество тромбоцитов", "MPV", "PLCR",
            "Спонтанная агрегация", "Индуц. агрегация 1 мкМоль АДФ", "Индуц. агрегация 5 мкМоль АДФ",
            "Индуц. агрегация 15 мкл арахидоновой кислоты", "Генотип CYP2C19", "Генотип ABCB1",
            "Препараты", "Состояние агрегации", "Скорость выведения клопидогрела (ABCB1)",
            "Модуль 1", "Модуль 2", "Модуль 3", "Коэффициент прогноза", "Оценка прогноза",
            # Новые поля для оценки риска ЖКК
            "Язвенная болезнь в анамнезе", "Желудочно-кишечное кровотечение в анамнезе",
            "Использование НПВП", "Прием ГКС", "Возраст ≥ 65 лет", "Диспепсия",
            "Желудочно-пищеводный рефлюкс", "Инфицирование H. pylori", "Хроническое употребление алкоголя",
            "Сумма баллов ЖКК"
        ]
        ws.append(headers)
    return wb, ws

def autofit_columns(ws):
    for column_cells in ws.columns:
        max_length = 0
        column = column_cells[0].column  # номер колонки (1, 2, 3...)
        for cell in column_cells:
            try:
                cell_length = len(str(cell.value))
                if cell_length > max_length:
                    max_length = cell_length
            except:
                pass
        adjusted_width = max_length + 2  # небольшой запас
        ws.column_dimensions[get_column_letter(column)].width = adjusted_width

def append_patient_data(filename, data_row):
    wb, ws = create_or_load_workbook(filename)
    if ws is not None:
        ws.append(data_row)
        autofit_columns(ws)
        wb.save(filename)
    else:
        raise ValueError("Не удалось создать или получить рабочий лист Excel")

def set_cell_border(cell, border_style="single", border_size=4, border_color="000000"):
    """Устанавливает границы для ячейки таблицы в docx"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Создаем элемент границы
    tcBorders = OxmlElement('w:tcBorders')
    
    # Устанавливаем границы для всех сторон
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), border_style)
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    
    tcPr.append(tcBorders)

def add_table_with_title(doc, headers, rows, title, column_widths=None):
    """Добавляет таблицу с заголовком, обеспечивая перенос на новую страницу при необходимости"""
    # Добавляем заголовок таблицы
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].bold = True
    title_para.runs[0].font.size = Pt(12)
    
    # Добавляем разрыв страницы перед таблицей, если она не помещается
    # Word автоматически перенесет таблицу на следующую страницу при необходимости
    doc.add_paragraph()  # Пустая строка
    
    # Создаем таблицу
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Устанавливаем ширины колонок если указаны
    if column_widths:
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
    
    # Заполняем заголовки
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
    
    # Заполняем данные
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_data)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell)
    
    # Добавляем пустую строку после таблицы
    doc.add_paragraph()

def add_simple_table(doc, headers, rows, column_widths=None):
    """Добавляет таблицу без заголовка"""
    # Создаем таблицу
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Устанавливаем ширины колонок если указаны
    if column_widths:
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
    
    # Заполняем заголовки
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
    
    # Заполняем данные
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_data)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell)
    
    # Добавляем пустую строку после таблицы
    doc.add_paragraph()

class ReportWindow(QWidget):
    def __init__(self, report_text, patient_data=None, excel_filename="patients.xlsx"):
        super().__init__()
        self.setWindowTitle("📋 Полный медицинский отчет по пациенту")
        self.resize(900, 700)
        self.patient_data = patient_data
        self.excel_filename = excel_filename
        self.current_report_data = None
        
        # Главный layout
        main_layout = QVBoxLayout(self)
        
        # Заголовок
        header_label = QLabel("МЕДИЦИНСКИЙ ОТЧЕТ ПО ПАЦИЕНТу")
        header_label.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #2c3e50;
                background-color: #ecf0f1;
                padding: 15px;
                border: 2px solid #3498db;
                border-radius: 8px;
                margin: 5px;
            }
        """)
        header_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(header_label)
        
        # Область прокрутки для текста
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Текстовое поле
        self.text = QTextEdit()
        self.text.setReadOnly(True)
        self.text.setStyleSheet("""
            QTextEdit {
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                font-size: 12px;
                line-height: 1.4;
                background-color: #f8f9fa;
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                padding: 15px;
                color: #2c3e50;
            }
            QTextEdit:focus {
                border: 2px solid #3498db;
            }
        """)
        
        # Форматируем текст отчета
        formatted_text = self.format_report_text(report_text)
        self.text.setText(formatted_text)
        
        scroll_area.setWidget(self.text)
        main_layout.addWidget(scroll_area)
        
        # Кнопки действий
        button_layout = QHBoxLayout()
        
        # Кнопка копирования
        copy_button = QPushButton("📋 Копировать отчет")
        copy_button.clicked.connect(self.copy_to_clipboard)
        copy_button.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                border: 2px solid #27ae60;
                border-radius: 5px;
                background-color: #27ae60;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:pressed {
                background-color: #1e8449;
            }
        """)
        
        # Кнопка сохранения в DOC
        doc_button = QPushButton("📝 Сохранить в DOC")
        doc_button.clicked.connect(self.save_to_doc)
        doc_button.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                border: 2px solid #3498db;
                border-radius: 5px;
                background-color: #3498db;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
        """)
        
        # Кнопка закрытия
        close_button = QPushButton("❌ Закрыть")
        close_button.clicked.connect(self.close)
        close_button.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                border: 2px solid #7f8c8d;
                border-radius: 5px;
                background-color: #7f8c8d;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #636e72;
            }
            QPushButton:pressed {
                background-color: #2d3436;
            }
        """)
        
        button_layout.addWidget(copy_button)
        button_layout.addWidget(doc_button)
        button_layout.addStretch()
        button_layout.addWidget(close_button)
        
        main_layout.addLayout(button_layout)
        
        # Стили для окна
        self.setStyleSheet("""
            QWidget {
                background-color: #ecf0f1;
            }
            QScrollArea {
                border: none;
                background-color: transparent;
            }
            QScrollBar:vertical {
                background-color: #f0f0f0;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #c0c0c0;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #a0a0a0;
            }
        """)
    
    def format_report_text(self, text):
        """Для HTML контента просто возвращаем как есть"""
        return text
    
    def copy_to_clipboard(self):
        """Копирует отчет в буфер обмена"""
        clipboard = QApplication.clipboard()
        clipboard.setText(self.text.toPlainText())
        QMessageBox.information(self, "Копирование", "Отчет скопирован в буфер обмена!")
    
    def save_to_doc(self):
        """Сохраняет отчет в DOC файл с табличным форматом"""
        if not self.current_report_data:
            QMessageBox.warning(self, "Предупреждение", "Нет данных для сохранения в DOC")
            return
            
        filename, _ = QFileDialog.getSaveFileName(
            self, 
            "Сохранить отчет в DOC", 
            f"медицинский_отчет_{QDate.currentDate().toString('yyyy-MM-dd')}.docx",
            "Word Documents (*.docx);;Все файлы (*)"
        )
        if filename:
            try:
                # Создаем новый документ
                doc = docx.Document()
                
                # Добавляем заголовок
                title = doc.add_heading('РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем основную информацию
                doc.add_paragraph(f"Дата обследования: {self.current_report_data['date']}")
                doc.add_paragraph(f"ФИО / № истории болезни: {self.current_report_data['name_or_record']}")
                doc.add_paragraph(f"Обследование: {self.current_report_data['examination_type']}")
                doc.add_paragraph(f"Возраст: {self.current_report_data['age']}")
                doc.add_paragraph()
                
                # Добавляем информацию о препаратах
                doc.add_paragraph().add_run("Прием антиагрегантов:").bold = True
                doc.add_paragraph(f"Антиагреганты, которые пациент принимает: {self.current_report_data['drugs']}")
                doc.add_paragraph()
                
                # Добавляем таблицы с данными
                add_table_with_title(doc, 
                    ["Параметр", "Результат пациента", "Критерий", "Оценка", "Прогноз"],
                    self.current_report_data['main_table_rows'],
                    "Прием антиагрегатов:"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента", "Оценка метаболизма", "Рекомендации"],
                    self.current_report_data['cyp_table_rows'],
                    "КОРРЕКЦИя ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ ГЕНОТИПА CYP 2C19"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента ABCB1", "Оценка транспорта", "Рекомендации"],
                    self.current_report_data['abcb1_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ АКТИВНОСТИ ТРАНСПОРТНОЙ СИСТЕМЫ P-ГЛИКОПРОТЕИНА"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['ticagrelor_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ ТИКАГРЕЛОРОМ"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 15 мкл арахидоновой кислоты, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['aspirin_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ АЦЕТИЛСАЛИЦИЛОВОЙ КИСЛОТОЙ"
                )
                
                # Добавляем таблицу прогноза непереносимости
                add_table_with_title(doc,
                    ["Параметр", "Значение", "Баллы"],
                    self.current_report_data['gi_bleeding_table_rows'],
                    "ПРОГНОЗ НЕПЕРЕНОСИМОСТИ ФАРМАКОТЕРАПИИ"
                )
                
                # Сохраняем документ
                doc.save(filename)
                QMessageBox.information(self, "Сохранение", f"Отчет сохранен в DOC файл:\n{filename}")
                
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить DOC файл:\n{str(e)}")
                print(f"Ошибка при сохранении DOC: {e}")

class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Антиагрегантная терапия")
        self.resize(1000, 800)
        
        # Создаем главный layout
        main_layout = QVBoxLayout(self)
        
        # Создаем область прокрутки
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Создаем контейнер для содержимого
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)

        # === ГРУППА 1: ОСНОВНЫЕ ДАННЫЕ ПАЦИЕНТА ===
        basic_group = QGroupBox("Основные данные пациента")
        basic_layout = QFormLayout()
        
        # Поля ввода
        self.date = QLineEdit()
        self.date.setPlaceholderText("Введите дату (дд.мм.гггг)")
        self.date.setText(QDate.currentDate().toString("dd.MM.yyyy"))
        basic_layout.addRow("Дата обследования:", self.date)

        self.name_or_record = QLineEdit()
        self.name_or_record.setPlaceholderText("Введите ФИО или номер истории болезни")
        basic_layout.addRow("ФИО / № истории болезни:", self.name_or_record)

        self.examination_type = QComboBox()
        self.examination_type.addItems(["Стационар", "Амбулаторно"])
        basic_layout.addRow("Обследование:", self.examination_type)
        
        # Поля выбора
        self.gender = QComboBox()
        self.gender.addItem("")  # Для необязательного выбора
        self.gender.addItems([g.value for g in Gender])
        basic_layout.addRow("Пол (выберите):", self.gender)
        
        # Поля ввода
        self.age = QLineEdit()
        self.age.setPlaceholderText("Введите возраст (лет)")
        basic_layout.addRow("Возраст:", self.age)

        self.weight = QLineEdit()
        self.weight.setPlaceholderText("Введите вес (кг)")
        basic_layout.addRow("Вес:", self.weight)

        self.height_field = QLineEdit()
        self.height_field.setPlaceholderText("Введите рост (см)")
        basic_layout.addRow("Рост:", self.height_field)
        
        basic_group.setLayout(basic_layout)
        layout.addWidget(basic_group)

        # === ГРУППА 2: БИОХИМИЧЕСКИЕ ПОКАЗАТЕЛИ ===
        bio_group = QGroupBox("Биохимические показатели")
        bio_layout = QFormLayout()
        
        self.creatinine = QLineEdit()
        self.creatinine.setPlaceholderText("Введите креатинин (мкмоль/л)")
        bio_layout.addRow("Креатинин:", self.creatinine)
        
        bio_group.setLayout(bio_layout)
        layout.addWidget(bio_group)

        # === ГРУППА 3: ТРОМБОЦИТАРНЫЕ ПОКАЗАТЕЛИ ===
        platelet_group = QGroupBox("Тромбоцитарные показатели")
        platelet_layout = QFormLayout()
        
        # Добавляем поле для количества тромбоцитов
        self.platelet_count = QLineEdit()
        self.platelet_count.setPlaceholderText("Введите количество тромбоцитов (×10⁹/л)")
        platelet_layout.addRow("Количество тромбоцитов, ×10⁹/л:", self.platelet_count)
        
        self.mpv = QLineEdit()
        self.mpv.setPlaceholderText("Введите MPV (фл)")
        platelet_layout.addRow("Величина тромбоцитов MPV:", self.mpv)

        self.plcr = QLineEdit()
        self.plcr.setPlaceholderText("Введите PLCR (%)")
        platelet_layout.addRow("Отн. кол-во больших тромбоцитов PLCR:", self.plcr)
        
        platelet_group.setLayout(platelet_layout)
        layout.addWidget(platelet_group)

        # === ГРУППА 4: АГРЕГАЦИя ТРОМБОЦИТОВ ===
        aggregation_group = QGroupBox("Агрегация тромбоцитов")
        aggregation_layout = QFormLayout()
        
        self.spontaneous_aggregation = QLineEdit()
        self.spontaneous_aggregation.setPlaceholderText("Введите спонтанную агрегацию (усл.ед.)")
        aggregation_layout.addRow("Спонтанная агрегация:", self.spontaneous_aggregation)

        self.induced_aggregation_1_ADP = QLineEdit()
        self.induced_aggregation_1_ADP.setPlaceholderText("Введите % агрегации")
        aggregation_layout.addRow("Индуц. агрегация 1 мкМоль АДФ:", self.induced_aggregation_1_ADP)

        self.induced_aggregation_5_ADP = QLineEdit()
        self.induced_aggregation_5_ADP.setPlaceholderText("Введите % агрегации")
        aggregation_layout.addRow("Индуц. агрегация 5 мкМоль АДФ:", self.induced_aggregation_5_ADP)

        self.induced_aggregation_15_ARA = QLineEdit()
        self.induced_aggregation_15_ARA.setPlaceholderText("Введите % агрегации")
        aggregation_layout.addRow("Индуц. агрегация 15 мкл арахидоновой кислоты:", self.induced_aggregation_15_ARA)
        
        aggregation_group.setLayout(aggregation_layout)
        layout.addWidget(aggregation_group)

        # === ГРУППА 5: ГЕНОТИПЫ ===
        genotype_group = QGroupBox("Генотипы")
        genotype_layout = QFormLayout()
        
        self.cyp2c19 = QComboBox()
        self.cyp2c19.addItem("")
        self.cyp2c19.addItems([c.value for c in CYP2C19])
        genotype_layout.addRow("Генотип CYP2C19:", self.cyp2c19)

        self.abcb1 = QComboBox()
        self.abcb1.addItem("")
        self.abcb1.addItems(["TT", "TC", "CC"])
        genotype_layout.addRow("Генотип ABCB1:", self.abcb1)
        
        genotype_group.setLayout(genotype_layout)
        layout.addWidget(genotype_group)

        # === ГРУППА 6: ОЦЕНКА РИСКА ЖЕЛУДОЧНО-КИШЕЧНОГО КРОВОТЕЧЕНИЯ ===
        gi_bleeding_group = QGroupBox("Оценка риска желудочно-кишечного кровотечения")
        gi_bleeding_layout = QFormLayout()
        
        # Поля для оценки риска ЖКК
        self.ulcer_history = QComboBox()
        self.ulcer_history.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Язвенная болезнь в анамнезе:", self.ulcer_history)

        self.gi_bleeding_history = QComboBox()
        self.gi_bleeding_history.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Желудочно-кишечное кровотечение в анамнезе:", self.gi_bleeding_history)

        self.nsaid_use = QComboBox()
        self.nsaid_use.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Использование НПВП:", self.nsaid_use)

        self.steroid_use = QComboBox()
        self.steroid_use.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Прием ГКС:", self.steroid_use)

        self.age_65 = QComboBox()
        self.age_65.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Возраст ≥ 65 лет:", self.age_65)

        self.dyspepsia = QComboBox()
        self.dyspepsia.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Диспепсия:", self.dyspepsia)

        self.gerd = QComboBox()
        self.gerd.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Желудочно-пищеводный рефлюкс:", self.gerd)

        self.h_pylori = QComboBox()
        self.h_pylori.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Инфицирование H. pylori:", self.h_pylori)

        self.alcohol_use = QComboBox()
        self.alcohol_use.addItems(["нет", "да"])
        gi_bleeding_layout.addRow("Хроническое употребление алкоголя:", self.alcohol_use)
        
        gi_bleeding_group.setLayout(gi_bleeding_layout)
        layout.addWidget(gi_bleeding_group)

        # === ГРУППА 7: ПРЕПАРАТЫ ===
        drugs_group = QGroupBox("Препараты")
        drugs_layout = QVBoxLayout()
        
        drugs_label = QLabel("Выберите принимаемый препарат:")
        drugs_layout.addWidget(drugs_label)
        
        # Создаем группу радиокнопок для выбора только одного препарата
        self.drugs_button_group = QButtonGroup(self)
        
        self.drug_aspirin = QRadioButton("АСК")
        self.drug_clopidogrel = QRadioButton("Клопидогрел")
        self.drug_aspirin_clopidogrel = QRadioButton("АСК+клопидогрел")
        self.drug_aspirin_ticagrelor = QRadioButton("АСК+тикагрелор")
        
        # Добавляем радиокнопки в группу
        self.drugs_button_group.addButton(self.drug_aspirin, 1)
        self.drugs_button_group.addButton(self.drug_clopidogrel, 2)
        self.drugs_button_group.addButton(self.drug_aspirin_clopidogrel, 3)
        self.drugs_button_group.addButton(self.drug_aspirin_ticagrelor, 4)
        
        # Устанавливаем "АСК" по умолчанию
        self.drug_aspirin.setChecked(True)
        
        drugs_layout.addWidget(self.drug_aspirin)
        drugs_layout.addWidget(self.drug_clopidogrel)
        drugs_layout.addWidget(self.drug_aspirin_clopidogrel)
        drugs_layout.addWidget(self.drug_aspirin_ticagrelor)
        
        drugs_group.setLayout(drugs_layout)
        layout.addWidget(drugs_group)

        # === ГРУППА 8: ДЕЙСТВИЯ ===
        actions_group = QGroupBox("⚙️ Действия")
        actions_layout = QVBoxLayout()
        
        self.report_button = QPushButton("📄 Сформировать полный отчет")
        self.report_button.clicked.connect(self.generate_report)
        actions_layout.addWidget(self.report_button)

        self.save_doc_button = QPushButton("📝 Сохранить отчет в DOC")
        self.save_doc_button.clicked.connect(self.save_report_to_doc)
        actions_layout.addWidget(self.save_doc_button)
        
        actions_group.setLayout(actions_layout)
        layout.addWidget(actions_group)

        # Устанавливаем контейнер в область прокрутки
        scroll_area.setWidget(content_widget)
        
        # Добавляем область прокрутки в главный layout
        main_layout.addWidget(scroll_area)

        self.excel_filename = DEFAULT_FILENAME
        self.patient_data = None
        self.current_report_html = ""
        self.current_report_data = None

        # Применяем стили
        self.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin-top: 1ex;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
            QLineEdit {
                padding: 5px;
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                background-color: #f8f9fa;
            }
            QLineEdit:focus {
                border: 2px solid #3498db;
                background-color: white;
            }
            QComboBox {
                padding: 5px;
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                background-color: #ecf0f1;
            }
            QComboBox:focus {
                border: 2px solid #3498db;
            }
            QPushButton {
                padding: 8px 16px;
                border: 2px solid #3498db;
                border-radius: 5px;
                background-color: #3498db;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
            QRadioButton {
                spacing: 8px;
                font-weight: normal;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
            QScrollArea {
                border: none;
                background-color: transparent;
            }
        """)

        # Подключение валидации к полям
        self.age.textChanged.connect(self.validate_age)
        self.weight.textChanged.connect(self.validate_weight)
        self.height_field.textChanged.connect(self.validate_height)
        self.creatinine.textChanged.connect(self.validate_creatinine)
        self.mpv.textChanged.connect(self.validate_mpv)
        self.plcr.textChanged.connect(self.validate_plcr)
        self.spontaneous_aggregation.textChanged.connect(self.validate_spontaneous_aggregation)
        self.induced_aggregation_1_ADP.textChanged.connect(self.validate_induced_aggregation_1_ADP)
        self.induced_aggregation_5_ADP.textChanged.connect(self.validate_induced_aggregation_5_ADP)
        self.induced_aggregation_15_ARA.textChanged.connect(self.validate_induced_aggregation_15_ARA)

    # Методы валидации
    def validate_age(self):
        try:
            age = int(self.age.text())
            if age <= 0 or age > 120:
                self.age.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.age.setStyleSheet("")
                return True
        except ValueError:
            self.age.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
            return False

    def validate_weight(self):
        try:
            weight = float(self.weight.text())
            if weight <= 0 or weight > 300:
                self.weight.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.weight.setStyleSheet("")
                return True
        except ValueError:
            self.weight.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
            return False

    def validate_height(self):
        try:
            height = float(self.height_field.text())
            if height <= 0 or height > 250:
                self.height_field.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.height_field.setStyleSheet("")
                return True
        except ValueError:
            if self.height_field.text():
                self.height_field.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.height_field.setStyleSheet("")
                return True

    def validate_creatinine(self):
        try:
            creatinine = float(self.creatinine.text())
            if creatinine <= 0 or creatinine > 1000:
                self.creatinine.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.creatinine.setStyleSheet("")
                return True
        except ValueError:
            if self.creatinine.text():
                self.creatinine.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.creatinine.setStyleSheet("")
                return True
    def validate_mpv(self):
        try:
            mpv = float(self.mpv.text())
            if mpv <= 0 or mpv > 20:
                self.mpv.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.mpv.setStyleSheet("")
                return True
        except ValueError:
            if self.mpv.text():
                self.mpv.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.mpv.setStyleSheet("")
                return True

    def validate_plcr(self):
        try:
            plcr = float(self.plcr.text())
            if plcr < 0 or plcr > 100:
                self.plcr.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.plcr.setStyleSheet("")
                return True
        except ValueError:
            if self.plcr.text():
                self.plcr.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.plcr.setStyleSheet("")
                return True

    def validate_spontaneous_aggregation(self):
        try:
            agg = float(self.spontaneous_aggregation.text())
            if agg < 0 or agg > 100:
                self.spontaneous_aggregation.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.spontaneous_aggregation.setStyleSheet("")
                return True
        except ValueError:
            if self.spontaneous_aggregation.text():
                self.spontaneous_aggregation.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.spontaneous_aggregation.setStyleSheet("")
                return True

    def validate_induced_aggregation_1_ADP(self):
        try:
            agg = float(self.induced_aggregation_1_ADP.text())
            if agg < 0 or agg > 100:
                self.induced_aggregation_1_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_1_ADP.setStyleSheet("")
                return True
        except ValueError:
            if self.induced_aggregation_1_ADP.text():
                self.induced_aggregation_1_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_1_ADP.setStyleSheet("")
                return True

    def validate_induced_aggregation_5_ADP(self):
        try:
            agg = float(self.induced_aggregation_5_ADP.text())
            if agg < 0 or agg > 100:
                self.induced_aggregation_5_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_5_ADP.setStyleSheet("")
                return True
        except ValueError:
            if self.induced_aggregation_5_ADP.text():
                self.induced_aggregation_5_ADP.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_5_ADP.setStyleSheet("")
                return True

    def validate_induced_aggregation_15_ARA(self):
        try:
            agg = float(self.induced_aggregation_15_ARA.text())
            if agg < 0 or agg > 100:
                self.induced_aggregation_15_ARA.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_15_ARA.setStyleSheet("")
                return True
        except ValueError:
            if self.induced_aggregation_15_ARA.text():
                self.induced_aggregation_15_ARA.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.induced_aggregation_15_ARA.setStyleSheet("")
                return True
    def validate_platelet_count(self):
        try:
            platelets = float(self.platelet_count.text())
            if platelets <= 0 or platelets > 1000:
                self.platelet_count.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.platelet_count.setStyleSheet("")
                return True
        except ValueError:
            if self.platelet_count.text():
                self.platelet_count.setStyleSheet("background-color: #ffcccc; border: 2px solid red;")
                return False
            else:
                self.platelet_count.setStyleSheet("")
                return True

    def validate_all_fields(self):
        validations = [
            self.validate_age(),
            self.validate_weight(),
            self.validate_height(),
            self.validate_creatinine(),
            self.validate_platelet_count(),  # Добавляем валидацию тромбоцитов
            self.validate_mpv(),
            self.validate_plcr(),
            self.validate_spontaneous_aggregation(),
            self.validate_induced_aggregation_1_ADP(),
            self.validate_induced_aggregation_5_ADP(),
            self.validate_induced_aggregation_15_ARA()
        ]

        if not all(validations):
            QMessageBox.warning(self, "Ошибка валидации",
                              "Пожалуйста, исправьте ошибки в полях (выделены красным)")
            return False
        return True

    def get_drug_cancellation_recommendation(self, platelet_count, drug_type):
        """Получить рекомендацию по отмене препарата на основе уровня тромбоцитов"""
        try:
            platelets = float(platelet_count)
            if platelets <= 10:
                if drug_type == "АСК":
                    return "Рекомендовано отменить АСК"
                elif drug_type == "АСК+тикагрелор":
                    return "Рекомендовано отменить тикагрелор и АСК"
                elif drug_type == "АСК+клопидогрел":
                    return "Рекомендовано отменить клопидогрел и АСК"
                elif drug_type == "клопидогрел":
                    return "Рекомендовано отменить клопидогрел"
            elif 10 < platelets <= 30:
                if drug_type in ["клопидогрел", "АСК+клопидогрел"]:
                    return "Рекомендовано отменить клопидогрел"
                elif drug_type in ["АСК+тикагрелор", "Тикагрелор"]:
                    return "Рекомендовано отменить тикагрелор"
                else:
                    return "Прием может быть продолжен"
            elif 30 < platelets <= 50: 
                if drug_type in ["АСК+тикагрелор", "Тикагрелор"]:
                    return "Рекомендовано отменить тикагрелор"
                elif drug_type in ["клопидогрел", "АСК", "АСК+клопидогрел"]:
                    return "Прием может быть продолжен"
                else: "Прием может быть продолжен"
            else:
                return "Прием может быть продолжен"
        except:
            return "Не определено"

    def format_html_table(self, headers, rows):
        html = '<table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 12px;">'
        
        # Заголовки
        html += '<tr style="background-color: #f2f2f2; font-weight: bold;">'
        for header in headers:
            html += f'<th style="border: 1px solid #000; padding: 8px; text-align: center;">{header}</th>'
        html += '</tr>'
        
        # Данные
        for row in rows:
            html += '<tr>'
            for cell in row:
                html += f'<td style="border: 1px solid #000; padding: 8px; text-align: center;">{cell}</td>'
            html += '</tr>'
        
        html += '</table>'
        return html

    def get_selected_drug(self):
        if self.drug_aspirin.isChecked():
            return "АСК"
        elif self.drug_clopidogrel.isChecked():
            return "клопидогрел"
        elif self.drug_aspirin_clopidogrel.isChecked():
            return "АСК+клопидогрел"
        elif self.drug_aspirin_ticagrelor.isChecked():
            return "АСК+тикагрелор"
        else:
            return ""

    def calculate_gi_bleeding_score(self):
        """Рассчитывает сумму баллов для оценки риска ЖКК"""
        score = 0
        
        # Преобразуем "да"/"нет" в 1/0
        if self.ulcer_history.currentText() == "да":
            score += 1
        if self.gi_bleeding_history.currentText() == "да":
            score += 1
        if self.nsaid_use.currentText() == "да":
            score += 1
        if self.steroid_use.currentText() == "да":
            score += 1
        if self.age_65.currentText() == "да":
            score += 1
        if self.dyspepsia.currentText() == "да":
            score += 1
        if self.gerd.currentText() == "да":
            score += 1
        if self.h_pylori.currentText() == "да":
            score += 1
        if self.alcohol_use.currentText() == "да":
            score += 1
            
        return score

    def save_report_to_doc(self):
        if not hasattr(self, 'current_report_data') or not self.current_report_data:
            QMessageBox.warning(self, "Предупреждение", "Сначала сформируйте отчет")
            return
            
        filename, _ = QFileDialog.getSaveFileName(
            self, 
            "Сохранить отчет в DOC", 
            f"медицинский_отчет_{QDate.currentDate().toString('yyyy-MM-dd')}.docx",
            "Word Documents (*.docx);;Все файлы (*)"
        )
        if filename:
            try:
                # Создаем новый документ
                doc = docx.Document()
                
                # Добавляем заголовок
                title = doc.add_heading('РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем основную информацию
                doc.add_paragraph(f"Дата обследования: {self.current_report_data['date']}")
                doc.add_paragraph(f"ФИО / № истории болезни: {self.current_report_data['name_or_record']}")
                doc.add_paragraph(f"Обследование: {self.current_report_data['examination_type']}")
                doc.add_paragraph(f"Возраст: {self.current_report_data['age']}")
                doc.add_paragraph()
                
                # Добавляем информацию о препаратах
                doc.add_paragraph().add_run("Прием антиагрегантов:").bold = True
                doc.add_paragraph(f"Антиагреганты, которые пациент принимает: {self.current_report_data['drugs']}")
                doc.add_paragraph()
                
                # Добавляем таблицы с данными
                add_table_with_title(doc, 
                    ["Параметр", "Результат пациента", "Критерий", "Оценка", "Прогноз"],
                    self.current_report_data['main_table_rows'],
                    "Прием антиагрегатов:"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента", "Оценка метаболизма", "Рекомендации"],
                    self.current_report_data['cyp_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ ГЕНОТИПА CYP 2C19"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Генотип пациента ABCB1", "Оценка транспорта", "Рекомендации"],
                    self.current_report_data['abcb1_table_rows'],
                    "КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ АКТИВНОСТИ ТРАНСПОРТНОЙ СИСТЕМЫ P-ГЛИКОПРОТЕИНА"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 5 мкМоль АДФ, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['ticagrelor_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ ТИКАГРЕЛОРОМ"
                )
                
                add_table_with_title(doc,
                    ["Индуцированная агрегация 15 мкл арахидоновой кислоты, % Т-крывая", "Критерий", "Состояние агрегации", "Рекомендации"],
                    self.current_report_data['aspirin_table_rows'],
                    "КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ АЦЕТИЛСАЛИЦИЛОВОЙ КИСЛОТОЙ"
                )
                
                # Добавляем таблицу прогноза непереносимости
                add_table_with_title(doc,
                    ["Параметр", "Значение", "Баллы"],
                    self.current_report_data['gi_bleeding_table_rows'],
                    "ПРОГНОЗ НЕПЕРЕНОСИМОСТИ ФАРМАКОТЕРАПИИ"
                )
                
                # Сохраняем документ
                doc.save(filename)
                QMessageBox.information(self, "Сохранение", f"Отчет сохранен в DOC файл:\n{filename}")
                
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить DOC файл:\n{str(e)}")
                print(f"Ошибка при сохранении DOC: {e}")

    def generate_report(self):
        try:
            if not self.validate_all_fields():
                return
            
            # Сбор данных
            date = self.date.text() if self.date.text() else QDate.currentDate().toString("dd.MM.yyyy")
            name_or_record = self.name_or_record.text() if self.name_or_record.text() else "____________________________________"
            age = int(self.age.text()) if self.age.text() else 0
            examination_type = self.examination_type.currentText()
            gender = self.gender.currentText()
            weight = float(self.weight.text()) if self.weight.text() else 0
            creatinine = float(self.creatinine.text()) if self.creatinine.text() else 0
            
            # Расчет КК и СКФ
            ccr = calculate_creatinine_clearance(age, weight, gender, creatinine)
            gfr = calculate_ckd_epi(age, gender, creatinine)
            
            # Получаем данные агрегации
            T_adp = float(self.induced_aggregation_5_ADP.text()) if self.induced_aggregation_5_ADP.text() else None
            T_ara = float(self.induced_aggregation_15_ARA.text()) if self.induced_aggregation_15_ARA.text() else None
            
            # Получаем количество тромбоцитов
            platelet_count = self.platelet_count.text() if self.platelet_count.text() else "______"
            
            # Получаем рекомендации по отмене препаратов
            selected_drug = self.get_selected_drug()
            drug_cancellation = self.get_drug_cancellation_recommendation(platelet_count, selected_drug)
            
            # Генетические данные
            cyp_genotype = self.cyp2c19.currentText() if self.cyp2c19.currentText() else "______"
            abcb1_genotype = self.abcb1.currentText() if self.abcb1.currentText() else "______"
            
            # Данные о терапии
            selected_drug = self.get_selected_drug()
            drugs_str = selected_drug if selected_drug else "___________"

            # Расчет коэффициента прогноза
            try:
                gender_val = 1 if self.gender.currentText() == "Муж" else 2 if self.gender.currentText() == "Жен" else 0
                prognosis_value = mod1(
                    gender_val,
                    float(self.age.text()) if self.age.text() else 0,
                    float(self.weight.text()) if self.weight.text() else 0,
                    float(self.height_field.text()) if self.height_field.text() else 0,
                    float(self.creatinine.text()) if self.creatinine.text() else 0,
                    0,  # Клиренс креатинина удален
                    float(self.mpv.text()) if self.mpv.text() else 0,
                    float(self.plcr.text()) if self.plcr.text() else 0,
                    float(self.spontaneous_aggregation.text()) if self.spontaneous_aggregation.text() else 0,
                    float(self.induced_aggregation_1_ADP.text()) if self.induced_aggregation_1_ADP.text() else 0,
                    float(self.induced_aggregation_5_ADP.text()) if self.induced_aggregation_5_ADP.text() else 0,
                    float(self.induced_aggregation_15_ARA.text()) if self.induced_aggregation_15_ARA.text() else 0
                )
                prognosis_evaluation = mod1_text(prognosis_value)
            except Exception as e:
                prognosis_value = "Ошибка расчета"
                prognosis_evaluation = ("Ошибка", ["Ошибка расчета коэффициента прогноза"])

            # Расчет оценки риска ЖКК
            gi_bleeding_score = self.calculate_gi_bleeding_score()

            # Сохраняем данные для DOC экспорта
            self.current_report_data = {
                'date': date,
                'name_or_record': name_or_record,
                'examination_type': examination_type,
                'age': age,
                'drugs': drugs_str,
                'main_table_rows': [],
                'cyp_table_rows': [],
                'abcb1_table_rows': [],
                'ticagrelor_table_rows': [],
                'aspirin_table_rows': [],
                'gi_bleeding_table_rows': []
            }

            # Формируем HTML отчет и данные для таблиц
            html_report = f"""
            <html>
            <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ text-align: center; font-size: 16px; font-weight: bold; margin-bottom: 20px; }}
                .section {{ margin: 20px 0; }}
                .section-title {{ font-size: 14px; font-weight: bold; margin-bottom: 10px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 8px; text-align: center; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
            </style>
            </head>
            <body>
                <div class="header">РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ</div>
                
                <p><strong>Дата обследования:</strong> {date}</p>
                <p><strong>ФИО / № истории болезни:</strong> {name_or_record}</p>
                <p><strong>Обследование:</strong> {examination_type}</p>
                <p><strong>Возраст:</strong> {age}</p>
                
                <div class="section">
                    <div class="section-title">Прием антиагрегантов:</div>
                    <p><strong>Антиагреганты, которые пациент принимает:</strong> {drugs_str}</p>
                </div>
            """

            # Таблица 1: Основные результаты
            main_table_headers = ["Параметр", "Результат пациента", "Критерий", "Оценка", "Прогноз"]
            main_table_rows = []

            # Строка 1: Коэффициент прогноза
            if isinstance(prognosis_value, (int, float)):
                if prognosis_value <= 1.56:
                    criterion = "≤ 1.56"
                    evaluation = "Благоприятная"
                    prognosis_text = "Неблагоприятных событий в течение года не ожидается"
                elif 1.561 <= prognosis_value <= 2.087:
                    criterion = "1.561-2.087"
                    evaluation = "Неблагоприятная"
                    prognosis_text = "Возможны обращения за медицинской помощью в течение ближайшего года"
                else:
                    criterion = "> 2.08"
                    evaluation = "Риск повторных сосудистых событий"
                    prognosis_text = "Высокий риск повторного инфаркта и летальный исход"
                main_table_rows.append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    f"{prognosis_value:.3f}",
                    criterion,
                    evaluation,
                    prognosis_text
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    f"{prognosis_value:.3f}",
                    criterion,
                    evaluation,
                    prognosis_text
                ])
            else:
                main_table_rows.append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    prognosis_value,
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Коэффициент прогноза неблагоприятных событий пациента с ОКС",
                    prognosis_value,
                    "-",
                    "-",
                    "-"
                ])

            # Строка 2: Индуцированная агрегация 5 мкМоль АДФ
            if T_adp is not None:
                if T_adp <= 10:
                    criterion_adp = "T ≤ 10 %"
                    evaluation_adp = "Агрегация тромбоцитов значительно подавлена"
                    prognosis_adp = "Риск геморрагических осложнений"
                elif 10 < T_adp and T_adp < 25:
                    criterion_adp = "10 - 25 %"
                    evaluation_adp = "Агрегация тромбоцитов умеренно подавлена"
                    prognosis_adp = "Терапия эффективна"
                else:
                    criterion_adp = "T ≥ 25 %"
                    evaluation_adp = "Агрегация тромбоцитов сохранена"
                    prognosis_adp = "Терапия неэффективна"

                main_table_rows.append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    f"{T_adp}%",
                    criterion_adp,
                    evaluation_adp,
                    prognosis_adp
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    f"{T_adp}%",
                    criterion_adp,
                    evaluation_adp,
                    prognosis_adp
                ])
            else:
                main_table_rows.append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    "______",
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                    "______",
                    "-",
                    "-",
                    "-"
                ])

            # Строка 3: Генотип CYP 2C19
            if cyp_genotype != "______":
                if cyp_genotype == "CYP 2c19*1":
                    evaluation_cyp = "Нормальный метаболизм клопидогрела"
                    prognosis_cyp = "Эффективность терапии клопидогрелом"
                elif cyp_genotype in ["CYP 2c19*2", "CYP 2c19*3"]:
                    evaluation_cyp = "Замедленный метаболизм клопидогрела"
                    prognosis_cyp = "Возможна резистентность к клопидогрелу"
                elif cyp_genotype == "CYP 2c19*17":
                    evaluation_cyp = "Ускоренный метаболизм клопидогрела"
                    prognosis_cyp = "Возможно угнетение агрегации, риск геморрагических осложнений"
                else:
                    evaluation_cyp = "Неизвестный генотип"
                    prognosis_cyp = "Требуется дополнительное исследование"

                main_table_rows.append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    cyp_genotype,
                    cyp_genotype,
                    evaluation_cyp,
                    prognosis_cyp
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    cyp_genotype,
                    cyp_genotype,
                    evaluation_cyp,
                    prognosis_cyp
                ])
            else:
                main_table_rows.append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    "______",
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип CYP 2C19, влияющий на метаболизм клопидогрела у пациента",
                    "______",
                    "-",
                    "-",
                    "-"
                ])

            # Строка 4: Генотип ABCB1
            if abcb1_genotype != "______":
                if abcb1_genotype == "TT":
                    evaluation_abcb1 = "Выведение клопидогрела ускорено"
                    prognosis_abcb1 = "Вероятна резистентность к клопидогрелу"
                elif abcb1_genotype == "TC":
                    evaluation_abcb1 = "Незначительное ускорение выведения клопидогрела"
                    prognosis_abcb1 = "Клинически незначимое влияние эффективность фармакотерапии"
                elif abcb1_genotype == "CC":
                    evaluation_abcb1 = "Нормальная скорость выведения клопидогрела"
                    prognosis_abcb1 = "Влияния на эффективность терапии клопидогрелом нет"
                else:
                    evaluation_abcb1 = "Неизвестный генотип"
                    prognosis_abcb1 = "Требуется дополнительное исследование"

                main_table_rows.append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    abcb1_genotype,
                    abcb1_genotype,
                    evaluation_abcb1,
                    prognosis_abcb1
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    abcb1_genotype,
                    abcb1_genotype,
                    evaluation_abcb1,
                    prognosis_abcb1
                ])
            else:
                main_table_rows.append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    "______",
                    "-",
                    "-",
                    "-"
                ])
                
                self.current_report_data['main_table_rows'].append([
                    "Генотип ABCB1, влияющий на транспорт клопидогрела",
                    "______",
                    "-",
                    "-",
                    "-"
                ])

            # Добавляем основную таблицу в отчет
            html_report += self.format_html_table(main_table_headers, main_table_rows)

            html_report += f"""
            <div class="section">
                <div class="section-title">РАСЧЕТНЫЕ ПОКАЗАТЕЛИ ФУНКЦИИ ПОЧЕК</div>
                <p><strong>Клиренс креатинина (КК):</strong> {ccr if ccr is not None else 'Не рассчитан'} мл/мин</p>
                <p><strong>Скорость клубочковой фильтрации (СКФ):</strong> {gfr if gfr is not None else 'Не рассчитана'} мл/мин</p>
            """

            # Таблица 2: Коррекция терапии клопидогрелом (CYP2C19)
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ ГЕНОТИПА CYP 2C19</div>
            """

            cyp_table_headers = [
                "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Генотип пациента",
                "Оценка метаболизма",
                "Рекомендации"
            ]
            
            cyp_table_rows = []
            if T_adp is not None and cyp_genotype != "______":
                # Правильное определение критерия на основе значения T_adp
                if T_adp <= 10:
                    criterion = "T ≤ 10 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                elif 10 < T_adp < 25:
                    criterion = "10 - 25 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                else:
                    criterion = "T ≥ 25 %"
                    state = "Агрегация тромбоцитов сохранена"
                
                # Получаем рекомендации из модуля 2
                mod2_result = mod2(T_adp, cyp_genotype)
                
                cyp_table_rows.append([
                    f"{T_adp}%",
                    criterion,
                    state,
                    cyp_genotype,
                    mod2_result[1] if len(mod2_result) > 1 else "-",
                    mod2_result[2] if len(mod2_result) > 2 else "-"
                ])
                self.current_report_data['cyp_table_rows'].append([
                    f"{T_adp}%",
                    criterion,
                    state,
                    cyp_genotype,
                    mod2_result[1] if len(mod2_result) > 1 else "-",
                    mod2_result[2] if len(mod2_result) > 2 else "-"
                ])
            else:
                cyp_table_rows.append(["______", "-", "-", "-", "-", "-"])
                self.current_report_data['cyp_table_rows'].append(["______", "-", "-", "-", "-", "-"])

            html_report += self.format_html_table(cyp_table_headers, cyp_table_rows)
            html_report += "</div>"

            # Таблица 3: Коррекция терапии клопидогрелом (ABCB1)
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ТЕРАПИИ КЛОПИДОГРЕЛОМ С УЧЕТОМ АКТИВНОСТИ ТРАНСПОРТНОЙ СИСТЕМЫ Р-ГЛИКОПРОТЕИНА</div>
            """

            abcb1_table_headers = [
                "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Генотип пациента ABCB1",
                "Оценка транспорта",
                "Рекомендации"
            ]
            
            abcb1_table_rows = []
            if T_adp is not None and abcb1_genotype != "______":
                # Правильное определение критерия на основе значения T_adp
                if T_adp <= 10:
                    criterion = "T ≤ 10 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                elif 10 < T_adp < 25:
                    criterion = "10 - 25 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                else:
                    criterion = "T ≥ 25 %"
                    state = "Агрегация тромбоцитов сохранена"
                
                # Получаем рекомендации из модуля 3
                mod3_result = mod3(T_adp, abcb1_genotype)
                
                abcb1_table_rows.append([
                    f"{T_adp}%",
                    criterion,
                    state,
                    abcb1_genotype,
                    mod3_result[1] if len(mod3_result) > 1 else "-",
                    mod3_result[2] if len(mod3_result) > 2 else "-"
                ])
                self.current_report_data['abcb1_table_rows'].append([
                    f"{T_adp}%",
                    criterion,
                    state,
                    abcb1_genotype,
                    mod3_result[1] if len(mod3_result) > 1 else "-",
                    mod3_result[2] if len(mod3_result) > 2 else "-"
                ])
            else:
                abcb1_table_rows.append(["______", "-", "-", "-", "-", "-"])
                self.current_report_data['abcb1_table_rows'].append(["______", "-", "-", "-", "-", "-"])

            html_report += self.format_html_table(abcb1_table_headers, abcb1_table_rows)
            html_report += "</div>"

            # Таблица 4: Коррекция терапии тикагрелором
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ ТИКАГРЕлОРОМ</div>
            """

            ticagrelor_table_headers = [
                "Индуцированная агрегация 5 мкМоль АДФ, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Рекомендации"
            ]
            
            ticagrelor_table_rows = []
            if T_adp is not None:
                # Правильное определение критерия на основе значения T_adp
                if T_adp <= 10:
                    criterion = "T ≤ 10 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                elif 10 < T_adp < 25:
                    criterion = "10 - 25 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                else:
                    criterion = "T ≥ 25 %"
                    state = "Агрегация тромбоцитов сохранена"
                
                # Получаем рекомендации из модуля 4
                mod4_result = mod4(T_adp)
                
                ticagrelor_table_rows.append([
                    f"{T_adp}%",
                    criterion,
                    state,
                    mod4_result[1] if len(mod4_result) > 1 else "-"
                ])
                self.current_report_data['ticagrelor_table_rows'].append([
                    f"{T_adp}%",
                    criterion,
                    state,
                    mod4_result[1] if len(mod4_result) > 1 else "-"
                ])
            else:
                ticagrelor_table_rows.append(["______", "-", "-", "-"])
                self.current_report_data['ticagrelor_table_rows'].append(["______", "-", "-", "-"])

            html_report += self.format_html_table(ticagrelor_table_headers, ticagrelor_table_rows)
            html_report += "</div>"

            # Таблица 5: Коррекция терапии ацетилсалициловой кислотой
            html_report += """
            <div class="section">
                <div class="section-title">КОРРЕКЦИЯ ФАРМАКОТЕРАПИИ АЦЕТИЛСАЛИЦИЛОВОЙ КИСЛОТОЙ</div>
            """

            aspirin_table_headers = [
                "Индуцированная агрегация 15 мкл арахидоновой кислоты, % Т-кривая",
                "Критерий",
                "Состояние агрегации",
                "Рекомендации"
            ]
            
            aspirin_table_rows = []
            if T_ara is not None:
                # Правильное определение критерия на основе значения T_ara
                if T_ara <= 2:
                    criterion = "Т ≤ 2 %"
                    state = "Агрегация тромбоцитов значительно подавлена"
                elif 2 < T_ara < 8:
                    criterion = "2 - 8 %"
                    state = "Агрегация тромбоцитов умеренно подавлена"
                else:
                    criterion = "Т ≥ 8 %"
                    state = "Агрегация тромбоцитов сохранена"
                
                # Получаем рекомендации из модуля 5
                mod5_result = mod5(T_ara)
                
                aspirin_table_rows.append([
                    f"{T_ara}%",
                    criterion,
                    state,
                    mod5_result[1] if len(mod5_result) > 1 else "-"
                ])
                self.current_report_data['aspirin_table_rows'].append([
                    f"{T_ara}%",
                    criterion,
                    state,
                    mod5_result[1] if len(mod5_result) > 1 else "-"
                ])
            else:
                aspirin_table_rows.append(["______", "-", "-", "-"])
                self.current_report_data['aspirin_table_rows'].append(["______", "-", "-", "-"])

            html_report += self.format_html_table(aspirin_table_headers, aspirin_table_rows)
            html_report += "</div>"

            # Таблица 6: Прогноз непереносимости фармакотерапии
            html_report += """
            <div class="section">
                <div class="section-title">ПРОГНОЗ НЕПЕРЕНОСИМОСТИ ФАРМАКОТЕРАПИИ</div>
            """

            gi_bleeding_table_headers = ["Параметр", "Значение", "Баллы"]
            gi_bleeding_table_rows = []
            
            # Данные для таблицы оценки риска ЖКК
            gi_bleeding_table_rows.append(["1. Оценка риска желудочно-кишечного кровотечения", "", ""])
            gi_bleeding_table_rows.append(["Язвенная болезнь в анамнезе", self.ulcer_history.currentText(), "1" if self.ulcer_history.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Желудочно-кишечное кровотечение в анамнезе", self.gi_bleeding_history.currentText(), "1" if self.gi_bleeding_history.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Использование НПВП", self.nsaid_use.currentText(), "1" if self.nsaid_use.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Прием ГКС", self.steroid_use.currentText(), "1" if self.steroid_use.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Возраст ≥ 65 лет", self.age_65.currentText(), "1" if self.age_65.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Диспепсия", self.dyspepsia.currentText(), "1" if self.dyspepsia.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Желудочно-пищеводный рефлюкс", self.gerd.currentText(), "1" if self.gerd.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Инфицирование H. pylori", self.h_pylori.currentText(), "1" if self.h_pylori.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Хроническое употребление алкоголя", self.alcohol_use.currentText(), "1" if self.alcohol_use.currentText() == "да" else "0"])
            gi_bleeding_table_rows.append(["Всего - баллов", f"{gi_bleeding_score}", f"{gi_bleeding_score}"])
            
            gi_bleeding_table_rows.append(["", "", ""])
            gi_bleeding_table_rows.append(["2. Отмена препарата при уровне тромбоцитов", "", ""])
            gi_bleeding_table_rows.append(["Количество тромбоцитов", f"{platelet_count}×10⁹/л", ""])
            gi_bleeding_table_rows.append(["Текущая терапия", selected_drug, ""])
            gi_bleeding_table_rows.append(["Рекомендация по отмене", drug_cancellation, ""])
            gi_bleeding_table_rows.append(["АСК - при уровне тромбоцитов", "≤10×10⁹/л", ""])
            gi_bleeding_table_rows.append(["Клопидогрел - при уровне тромбоцитов", "≤30×10⁹/л", ""])
            gi_bleeding_table_rows.append(["Тикагрелор - при уровне тромбоцитов", "≤50×10⁹/л", ""])

            # Сохраняем данные для DOC
            self.current_report_data['gi_bleeding_table_rows'] = gi_bleeding_table_rows

            html_report += self.format_html_table(gi_bleeding_table_headers, gi_bleeding_table_rows)
            html_report += "</div>"

            # Закрываем HTML
            html_report += """
            </body>
            </html>
            """

            # Сохраняем текущий отчет
            self.current_report_html = html_report

            # Показываем отчет в диалоговом окне
            report_dialog = QDialog(self)
            report_dialog.setWindowTitle("Медицинский отчет")
            report_dialog.resize(900, 700)

            layout = QVBoxLayout(report_dialog)
            text_edit = QTextEdit()
            text_edit.setHtml(html_report)
            text_edit.setReadOnly(True)
            layout.addWidget(text_edit)

            close_button = QPushButton("Закрыть")
            close_button.clicked.connect(report_dialog.accept)
            layout.addWidget(close_button)

            report_dialog.exec()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка генерации отчета", 
                            f"Произошла ошибка при формировании отчета:\n{str(e)}")
            print(f"Ошибка в generate_report: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
